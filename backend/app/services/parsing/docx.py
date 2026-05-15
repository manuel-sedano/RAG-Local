"""Extracción de texto desde DOCX con python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from app.services.parsing.errors import FatalParserError, RecoverableParserError
from app.services.parsing.types import PageText, ParsedDocument


def _paragraph_section(paragraph: Paragraph) -> str | None:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading" in style_name or style_name.startswith("título"):
        return paragraph.text.strip() or None
    return None


def extract_docx(path: Path) -> ParsedDocument:
    try:
        document = DocxDocument(path)
    except Exception as e:
        raise FatalParserError("docx_open_failed", f"No se pudo abrir el DOCX: {e}") from e

    try:
        pages: list[PageText] = []
        blocks: list[str] = []
        current_section: str | None = None
        block_index = 0

        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            section = _paragraph_section(paragraph) or current_section
            if _paragraph_section(paragraph):
                current_section = section
            block_index += 1
            pages.append(PageText(page_number=block_index, text=text, section=section))
            prefix = f"[{section}] " if section else ""
            blocks.append(f"{prefix}{text}")

        full_text = "\n\n".join(blocks)
        page_count = max(len(pages), 1) if full_text else 0
        if not full_text.strip():
            raise RecoverableParserError(
                "docx_empty",
                "El DOCX no contiene texto extraíble.",
            )

        docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return ParsedDocument(
            mime_type=docx_mime,
            pages=pages,
            full_text=full_text,
            page_count=page_count,
            parser_used="python-docx",
            needs_ocr=False,
            metadata={"paragraph_count": len(pages)},
        )
    except (RecoverableParserError, FatalParserError):
        raise
    except Exception as e:
        raise RecoverableParserError("docx_extract_failed", f"Error al extraer DOCX: {e}") from e
