"""Tests unitarios de extractores PDF/DOCX/TXT y normalización."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument

from app.core.config import Settings, clear_settings_cache
from app.services.parsing.docx import extract_docx
from app.services.parsing.errors import FatalParserError, RecoverableParserError
from app.services.parsing.normalize import normalize_whitespace
from app.services.parsing.orchestrator import parse_document_file
from app.services.parsing.pdf import extract_pdf
from app.services.parsing.txt import extract_txt


def _settings(**overrides) -> Settings:
    clear_settings_cache()
    return Settings(environment="test", **overrides)


def _write_pdf(path: Path, pages: list[str]) -> None:
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def _write_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Introducción", level=1)
    doc.add_paragraph("Párrafo con viñeta.", style="List Bullet")
    doc.add_paragraph("Segunda viñeta.", style="List Bullet")
    doc.add_paragraph("Texto normal después de la lista.")
    doc.save(path)


def test_pdf_extract_text_and_page_count(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    _write_pdf(pdf_path, ["Página uno con texto.", "Página dos con más texto."])
    parsed = extract_pdf(pdf_path, ocr_min_chars_per_page=10)
    assert parsed.page_count == 2
    assert "Página uno" in parsed.full_text
    assert "Página dos" in parsed.full_text
    assert parsed.parser_used == "pymupdf"
    assert parsed.needs_ocr is False


def test_pdf_detects_low_text_for_ocr(tmp_path: Path) -> None:
    pdf_path = tmp_path / "scan.pdf"
    doc = fitz.open()
    doc.new_page()  # página vacía
    doc.save(pdf_path)
    doc.close()
    parsed = extract_pdf(pdf_path, ocr_min_chars_per_page=40)
    assert parsed.needs_ocr is True


def test_pdf_accented_characters(tmp_path: Path) -> None:
    pdf_path = tmp_path / "accents.pdf"
    _write_pdf(pdf_path, ["Información técnica: ñáéíóú — año 2024."])
    parsed = extract_pdf(pdf_path, ocr_min_chars_per_page=5)
    assert "ñáéíóú" in parsed.full_text or "Información" in parsed.full_text


def test_docx_extract_headings_and_lists(tmp_path: Path) -> None:
    docx_path = tmp_path / "lists.docx"
    _write_docx(docx_path)
    parsed = extract_docx(docx_path)
    assert parsed.page_count >= 1
    assert "Introducción" in parsed.full_text or "Párrafo" in parsed.full_text
    assert parsed.parser_used == "python-docx"


def test_txt_latin1_encoding(tmp_path: Path) -> None:
    txt_path = tmp_path / "latin1.txt"
    txt_path.write_bytes(b"Ma\xf1ana caf\xe9 a\xf1o")
    parsed = extract_txt(txt_path)
    assert parsed.encoding == "latin-1"
    assert "añana" in parsed.full_text or "Mañana" in parsed.full_text
    assert parsed.page_count == 1


def test_txt_utf8_normalization(tmp_path: Path) -> None:
    txt_path = tmp_path / "utf8.txt"
    txt_path.write_text("Línea uno\r\nLínea dos", encoding="utf-8")
    parsed = extract_txt(txt_path)
    assert "\r" not in parsed.full_text


def test_normalize_whitespace_hyphen_and_blanks() -> None:
    raw = "palabra-\ncontinuación\n\n\n\notra"
    out = normalize_whitespace(raw)
    assert "palabracontinuación" in out or "palabra-continuación" in out.replace("\n", "")
    assert "\n\n\n" not in out


def test_parse_document_file_integration(tmp_path: Path) -> None:
    pdf_path = tmp_path / "full.pdf"
    _write_pdf(pdf_path, ["Documento completo para orquestador."])
    settings = _settings(parse_timeout_seconds=30.0, parser_save_artifacts=False)
    parsed = parse_document_file(pdf_path, "application/pdf", settings)
    assert parsed.full_text
    assert parsed.parser_used == "pymupdf"


def test_parse_empty_pdf_raises_fatal(tmp_path: Path) -> None:
    pdf_path = tmp_path / "empty.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(pdf_path)
    doc.close()
    settings = _settings(ocr_min_chars_per_page=0, unstructured_enabled=False)
    with pytest.raises(FatalParserError) as exc:
        parse_document_file(pdf_path, "application/pdf", settings)
    assert exc.value.code == "empty_document"


def test_parse_timeout(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    pdf_path = tmp_path / "slow.pdf"
    _write_pdf(pdf_path, ["ok"])
    settings = _settings(parse_timeout_seconds=1.0)

    def _slow_extract(*_a, **_k):
        import time

        time.sleep(2)
        return extract_pdf(pdf_path, ocr_min_chars_per_page=10)

    monkeypatch.setattr(
        "app.services.parsing.orchestrator._extract_raw",
        lambda *_a, **_k: _slow_extract(),
    )
    with pytest.raises(RecoverableParserError) as exc:
        parse_document_file(pdf_path, "application/pdf", settings)
    assert exc.value.code == "parse_timeout"
